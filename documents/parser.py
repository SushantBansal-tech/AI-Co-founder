import csv
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader


class UnsupportedDocumentError(ValueError):
    pass


def parse_txt(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="ignore")


def parse_pdf(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    pages: list[str] = []

    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""

        if text.strip():
            pages.append(
                f"\n[PAGE {page_number}]\n{text.strip()}"
            )

    return "\n".join(pages)


def parse_docx(file_path: Path) -> str:
    document = DocxDocument(str(file_path))

    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    # Include Word tables as text.
    for table_index, table in enumerate(document.tables, start=1):
        paragraphs.append(f"\n[TABLE {table_index}]")

        for row in table.rows:
            values = [
                cell.text.strip().replace("\n", " ")
                for cell in row.cells
            ]
            paragraphs.append(" | ".join(values))

    return "\n".join(paragraphs)


def parse_csv(file_path: Path) -> str:
    rows: list[str] = []

    with file_path.open(
        "r",
        encoding="utf-8-sig",
        errors="ignore",
        newline="",
    ) as csv_file:
        reader = csv.reader(csv_file)

        for row in reader:
            rows.append(" | ".join(str(value) for value in row))

    return "\n".join(rows)


def parse_excel(file_path: Path) -> str:
    workbook = load_workbook(
        filename=str(file_path),
        read_only=True,
        data_only=True,
    )

    output: list[str] = []

    for sheet in workbook.worksheets:
        output.append(f"\n[SHEET: {sheet.title}]")

        for row in sheet.iter_rows(values_only=True):
            values = [
                "" if value is None else str(value)
                for value in row
            ]

            if any(value.strip() for value in values):
                output.append(" | ".join(values))

    return "\n".join(output)


def parse_document(file_path: str) -> str:
    path = Path(file_path)
    extension = path.suffix.lower()

    parsers = {
        ".txt": parse_txt,
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".csv": parse_csv,
        ".xlsx": parse_excel,
    }

    parser = parsers.get(extension)

    if parser is None:
        raise UnsupportedDocumentError(
            f"Unsupported file type: {extension}. "
            "Supported types: PDF, DOCX, TXT, CSV and XLSX."
        )

    text = parser(path)

    if not text.strip():
        raise ValueError(
            "No text could be extracted. "
            "The document may be empty or scanned and may require OCR."
        )

    return text